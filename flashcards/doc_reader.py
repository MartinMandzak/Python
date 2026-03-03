import pandas as pd
import numpy as np
import os
from docx import Document

# read through all the files in ./data
def read_file(path):
    file = Document(path)
    if not file.tables:
        return None
    raw_data = file.tables[0]
    final_data = []
    for row in raw_data.rows:
        final_data.append([cell.text.strip() for cell in row.cells])
    df = pd.DataFrame(final_data)
    df.replace('',np.nan, inplace=True)
    df.dropna(subset=[0], inplace=True)
    df.dropna(how='all', inplace=True)
    df.reset_index(drop=True, inplace=True)

    
    return df

# dir
path = './data/'
all_data = []
def read_all_in_dir(path):
    for filename in os.listdir(path):
        if filename.endswith(".docx"):
            df = read_file(os.path.join(path, filename))
            if df is not None:
                df['Source_File'] = filename
                all_data.append(df)

    master_df = pd.concat(all_data, ignore_index=True)
    return master_df


